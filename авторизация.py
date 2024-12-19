import sys  # Импортируем модуль sys для получения доступа к параметрам командной строки и другим функциям
import sqlite3  # Импортируем библиотеку для работы с SQLite базами данных
from PyQt6 import QtWidgets  # Импортируем основные компоненты для создания GUI из PyQt6
from PyQt6.QtWidgets import (  # Импортируем необходимые виджеты из модуля QtWidgets
    QMessageBox, QVBoxLayout, QTableWidget, QTableWidgetItem,
    QPushButton, QLineEdit, QDialog, QFormLayout, QDateTimeEdit,
    QLabel, QFileDialog
)
from PyQt6.QtCore import QDateTime, Qt  # Импортируем типы данных и флаги из QtCore
from docx import Document  # Импортируем класс Document для работы с файлами .docx


class AuthApp(QtWidgets.QWidget):  # Определяем класс приложения авторизации, наследующий от QWidget
    def __init__(self):  # Конструктор класса
        super().__init__()  # Вызываем конструктор родительского класса

        self.setWindowTitle("BOX BOX")  # Устанавливаем заголовок окна
        self.setGeometry(100, 100, 400, 250)  # Устанавливаем размеры и положение окна

        # Стилизация интерфейса
        self.setStyleSheet("""  
            QWidget {
                background-color: #D6EAF8;  
                color: #2C3E50;  
                font-family: Arial;  
                font-size: 14px; 
            }
            QLabel {
                color: #2C3E50;  
                font-weight: bold;  
            }
            QLineEdit {
                background-color: #FFFFFF;  
                border: 1px solid #000000;  
                border-radius: 5px;  
                padding: 5px;  
                color: #2C3E50;  
            }
            QPushButton {
                background-color: #3498DB;  
                border: none;  
                border-radius: 5px;  
                padding: 8px 15px;  
                color: #ECF0F1;  
                font-weight: bold;  
            }
            QPushButton:hover {
                background-color: #2980B9;  
            }
            QPushButton:pressed {
                background-color: #1C5980;  
            }
        """)

        # Создание виджетов
        self.username_label = QtWidgets.QLabel("Имя пользователя:", self)  # Метка для поля имени пользователя
        self.username_label.move(20, 50)  # Устанавливаем позицию метки

        self.username_input = QLineEdit(self)  # Поле ввода для имени пользователя
        self.username_input.move(160, 50)  # Устанавливаем позицию поля
        self.username_input.resize(200, 25)  # Устанавливаем размеры поля

        self.password_label = QtWidgets.QLabel("Пароль:", self)  # Метка для поля пароля
        self.password_label.move(20, 100)  # Устанавливаем позицию метки

        self.password_input = QLineEdit(self)  # Поле ввода для пароля
        self.password_input.setEchoMode(QLineEdit.EchoMode.Password)  # Устанавливаем режим скрытия символов
        self.password_input.move(160, 100)  # Устанавливаем позицию поля
        self.password_input.resize(200, 25)  # Устанавливаем размеры поля

        self.login_button = QPushButton("Войти", self)  # Кнопка входа
        self.login_button.move(150, 160)  # Устанавливаем позицию кнопки
        self.login_button.clicked.connect(self.check_credentials)  # Подключаем обработчик события нажатия

    def check_credentials(self):  # Метод для проверки учетных данных пользователя
        username = self.username_input.text()  # Получаем текст из поля имени пользователя
        password = self.password_input.text()  # Получаем текст из поля пароля

        try:
            conn = sqlite3.connect('myydatabase.db')  # Подключаемся к базе данных
            cursor = conn.cursor()  # Создаем курсор для выполнения запросов

            # Проверка логина и пароля
            cursor.execute("SELECT role FROM users WHERE username=? AND password=?",
                           (username, password))  # Выполняем SQL-запрос для проверки логина и пароля
            user = cursor.fetchone()  # Получаем результаты запроса

            if user:  # Если пользователь найден
                role = user[0]  # Получаем роль пользователя
                QMessageBox.information(self, "Успех", "Вы успешно вошли в систему!")  # Сообщение об успешном входе
                if role == 'admin':  # Если роль администратора
                    self.open_admin_window()  # Открываем окно администратора
                elif role == 'manager':  # Если роль менеджера
                    self.open_manager_window()  # Открываем окно менеджера
            else:  # Если пользователь не найден
                QMessageBox.warning(self, "Ошибка", "Неверное имя пользователя или пароль!")  # Сообщение об ошибке

        except Exception as e:  # Обработка ошибок
            QMessageBox.critical(self, "Ошибка", f"Произошла ошибка: {e}")  # Сообщение об ошибке

        finally:
            conn.close()  # Закрываем подключение к базе данных

    def open_admin_window(self):  # Метод для открытия окна администратора
        self.admin_window = AdminWindow()  # Создаем экземпляр окна администратора
        self.admin_window.show()  # Показываем его
        self.close()  # Закрываем текущее окно

    def open_manager_window(self):  # Метод для открытия окна менеджера
        self.manager_window = ManagerWindow()  # Создаем экземпляр окна менеджера
        self.manager_window.show()  # Показываем его
        self.close()  # Закрываем текущее окно


class AddClientDialog(QDialog):  # Класс диалогового окна для добавления клиента
    """Диалоговое окно для добавления клиента."""

    def __init__(self):  # Конструктор класса
        super().__init__()  # Вызываем конструктор родительского класса
        self.setWindowTitle("Добавить клиента")  # Устанавливаем заголовок окна
        self.setGeometry(100, 100, 400, 300)  # Устанавливаем размеры и положение окна

        layout = QFormLayout()  # Создаем компоновку форм

        # Создаем поля для ввода данных клиента
        self.first_name_input = QLineEdit()  # Поле ввода имени
        self.last_name_input = QLineEdit()  # Поле ввода фамилии
        self.phone_number_input = QLineEdit()  # Поле ввода номера телефона
        self.email_input = QLineEdit()  # Поле ввода Email'

        # Добавляем строки в компоновку
        layout.addRow("Имя:", self.first_name_input)
        layout.addRow("Фамилия:", self.last_name_input)
        layout.addRow("Телефон:", self.phone_number_input)
        layout.addRow("Email:", self.email_input)

        self.submit_button = QPushButton("Добавить")  # Кнопка для добавления клиента
        self.submit_button.clicked.connect(self.add_client_to_db)  # Подключаем обработчик нажатия
        layout.addWidget(self.submit_button)  # Добавляем кнопку в компоновку

        self.setLayout(layout)  # Устанавливаем компоновку для диалога

    def add_client_to_db(self):  # Метод для добавления клиента в базу данных
        """Добавляет клиента в базу данных."""
        first_name = self.first_name_input.text()  # Получаем имя клиента
        last_name = self.last_name_input.text()  # Получаем фамилию клиента
        phone_number = self.phone_number_input.text()  # Получаем номер телефона клиента
        email = self.email_input.text()  # Получаем email клиента

        if not (first_name and last_name and phone_number and email):  # Проверяем, что поля не пустые
            QMessageBox.warning(self, "Ошибка", "Заполните все поля!")  # Сообщение об ошибке
            return

        conn = sqlite3.connect('myydatabase.db')  # Подключаемся к базе данных
        cursor = conn.cursor()  # Создаем курсор

        # Выполняем SQL-запрос для добавления клиента
        cursor.execute(
            "INSERT INTO clients (first_name, last_name, phone_number, email) VALUES (?, ?, ?, ?)",
            (first_name, last_name, phone_number, email)
        )
        conn.commit()  # Подтверждаем изменения
        conn.close()  # Закрываем соединение

        QMessageBox.information(self, "Успех", "Клиент успешно добавлен!")  # Сообщение об успешном добавлении
        self.close()  # Закрываем диалоговое окно


class AddCarDialog(QDialog):  # Класс диалогового окна для добавления автомобиля
    """Диалоговое окно для добавления автомобиля."""

    def __init__(self):  # Конструктор класса
        super().__init__()  # Вызываем конструктор родительского класса
        self.setWindowTitle("Добавить автомобиль")  # Устанавливаем заголовок окна
        self.setGeometry(100, 100, 400, 300)  # Устанавливаем размеры и положение окна

        layout = QFormLayout()  # Создаем компоновку форм

        # Создаем поля для ввода данных автомобиля
        self.client_id_input = QLineEdit()  # Поле ввода ID клиента
        self.brand_input = QLineEdit()  # Поле ввода марки автомобиля
        self.model_input = QLineEdit()  # Поле ввода модели автомобиля
        self.car_number_input = QLineEdit()  # Поле ввода номера автомобиля

        # Добавляем строки в компоновку
        layout.addRow("ID клиента:", self.client_id_input)
        layout.addRow("Марка:", self.brand_input)
        layout.addRow("Модель:", self.model_input)
        layout.addRow("Номер авто:", self.car_number_input)

        self.submit_button = QPushButton("Добавить")  # Кнопка для добавления автомобиля
        self.submit_button.clicked.connect(self.add_car_to_db)  # Подключаем обработчик нажатия
        layout.addWidget(self.submit_button)  # Добавляем кнопку в компоновку

        self.setLayout(layout)  # Устанавливаем компоновку для диалога

    def add_car_to_db(self):  # Метод для добавления автомобиля в базу данных
        """Добавляет автомобиль в базу данных."""
        client_id = self.client_id_input.text()  # Получаем ID клиента
        brand = self.brand_input.text()  # Получаем марку автомобиля
        model = self.model_input.text()  # Получаем модель автомобиля
        car_number = self.car_number_input.text()  # Получаем номер автомобиля

        if not (client_id and brand and model and car_number):  # Проверяем, что поля не пустые
            QMessageBox.warning(self, "Ошибка", "Заполните все поля!")  # Сообщение об ошибке
            return

        conn = sqlite3.connect('myydatabase.db')  # Подключаемся к базе данных
        cursor = conn.cursor()  # Создаем курсор

        # Выполняем SQL-запрос для добавления автомобиля
        cursor.execute(
            "INSERT INTO cars (id_client, brand_car, car_model, car_number) VALUES (?, ?, ?, ?)",
            (client_id, brand, model, car_number)
        )
        conn.commit()  # Подтверждаем изменения
        conn.close()  # Закрываем соединение

        QMessageBox.information(self, "Успех", "Автомобиль успешно добавлен!")  # Сообщение об успешном добавлении
        self.close()  # Закрываем диалоговое окно


class AddOrderDialog(QDialog):  # Класс диалогового окна для создания заказа
    """Диалоговое окно для создания заказа."""

    def __init__(self):  # Конструктор класса
        super().__init__()  # Вызываем конструктор родительского класса
        self.setWindowTitle("Создать заказ")  # Устанавливаем заголовок окна
        self.setGeometry(100, 100, 400, 300)  # Устанавливаем размеры и положение окна

        layout = QFormLayout()  # Создаем компоновку форм

        # Создаем поля для ввода данных заказа
        self.car_id_input = QLineEdit()  # Поле ввода ID автомобиля
        self.user_id_input = QLineEdit()  # Поле ввода ID пользователя
        self.order_description_input = QLineEdit()  # Поле ввода описания заказа
        self.price_input = QLineEdit()  # Поле ввода цены
        self.date_input = QDateTimeEdit(QDateTime.currentDateTime())  # Поле выбора даты заказа с текущей датой

        # Добавляем строки в компоновку
        layout.addRow("ID автомобиля:", self.car_id_input)
        layout.addRow("ID пользователя:", self.user_id_input)
        layout.addRow("Описание заказа:", self.order_description_input)
        layout.addRow("Цена:", self.price_input)
        layout.addRow("Дата заказа:", self.date_input)

        self.submit_button = QPushButton("Создать")  # Кнопка для создания заказа
        self.submit_button.clicked.connect(self.add_order_to_db)  # Подключаем обработчик нажатия
        layout.addWidget(self.submit_button)  # Добавляем кнопку в компоновку

        self.setLayout(layout)  # Устанавливаем компоновку для диалога

    def add_order_to_db(self):  # Метод для добавления заказа в базу данных
        """Добавляет заказ в базу данных."""
        car_id = self.car_id_input.text()  # Получаем ID автомобиля
        user_id = self.user_id_input.text()  # Получаем ID пользователя
        description = self.order_description_input.text()  # Получаем описание заказа
        price = self.price_input.text()  # Получаем цену заказа
        order_date = self.date_input.dateTime().toString("yyyy-MM-dd HH:mm:ss")  # Получаем дату заказа в нужном формате

        if not (car_id and user_id and description and price):  # Проверяем, что поля не пустые
            QMessageBox.warning(self, "Ошибка", "Заполните все поля!")  # Сообщение об ошибке
            return

        conn = sqlite3.connect('myydatabase.db')  # Подключаемся к базе данных
        cursor = conn.cursor()  # Создаем курсор

        # Выполняем SQL-запрос для добавления заказа
        cursor.execute(
            "INSERT INTO orders (id_car, id_user, order_client, price, order_date) VALUES (?, ?, ?, ?, ?)",
            (car_id, user_id, description, price, order_date)
        )
        conn.commit()  # Подтверждаем изменения
        conn.close()  # Закрываем соединение

        QMessageBox.information(self, "Успех", "Заказ успешно создан!")  # Сообщение об успешном создании заказа
        self.close()  # Закрываем диалоговое окно


class AdminWindow(QtWidgets.QWidget):  # Класс окна администратора
    """Окно администратора для работы с клиентами, машинами и заказами."""

    def __init__(self):  # Конструктор класса
        super().__init__()  # Вызываем конструктор родительского класса
        self.setWindowTitle("Администратор - Управление данными")  # Устанавливаем заголовок окна
        self.setGeometry(100, 100, 600, 400)  # Устанавливаем размеры и положение окна

        # Стилизация окна
        self.setStyleSheet("""  
                   QWidget {
                       background-color: #D6EAF8;  
                       color: #2C3E50;  
                       font-family: Arial;  
                       font-size: 14px;  
                   }
                   QPushButton {
                       background-color: #3498DB;  
                       border: none;  
                       border-radius: 5px;  
                       padding: 10px 15px;  
                       color: #ECF0F1;  
                       font-weight: bold;  
                   }
                   QPushButton:hover {
                       background-color: #2980B9;  
                   }
                   QPushButton:pressed {
                       background-color: #1C5980;  
                   }
               """)

        layout = QVBoxLayout()  # Создаем вертикальную компоновку

        # Заголовок панели администратора
        self.title_label = QLabel("Панель администратора")  # Создаем метку заголовка
        self.title_label.setAlignment(Qt.AlignmentFlag.AlignCenter)  # Центрируем заголовок
        self.title_label.setStyleSheet(
            "font-size: 18px; font-weight: bold; margin-bottom: 20px;")  # Стили для заголовка
        layout.addWidget(self.title_label)  # Добавляем заголовок в компоновку

        # Кнопки для добавления данных
        self.add_client_button = QPushButton("Добавить клиента")  # Кнопка для добавления клиента
        self.add_client_button.clicked.connect(self.add_client)  # Подключаем обработчик нажатия
        layout.addWidget(self.add_client_button)  # Добавляем кнопку в компоновку

        self.add_car_button = QPushButton("Добавить автомобиль")  # Кнопка для добавления автомобиля
        self.add_car_button.clicked.connect(self.add_car)  # Подключаем обработчик нажатия
        layout.addWidget(self.add_car_button)  # Добавляем кнопку в компоновку

        self.add_order_button = QPushButton("Создать заказ")  # Кнопка для создания заказа
        self.add_order_button.clicked.connect(self.add_order)  # Подключаем обработчик нажатия
        layout.addWidget(self.add_order_button)  # Добавляем кнопку в компоновку

        self.setLayout(layout)  # Устанавливаем компоновку для окна

    def add_client(self):  # Метод для открытия диалога добавления клиента
        dialog = AddClientDialog()  # Создаем экземпляр диалога
        dialog.exec()  # Показываем диалоговое окно

    def add_car(self):  # Метод для открытия диалога добавления автомобиля
        dialog = AddCarDialog()  # Создаем экземпляр диалога
        dialog.exec()  # Показываем диалоговое окно

    def add_order(self):  # Метод для открытия диалога создания заказа
        dialog = AddOrderDialog()  # Создаем экземпляр диалога
        dialog.exec()  # Показываем диалоговое окно


class ManagerWindow(QtWidgets.QWidget):  # Класс окна менеджера
    """Окно менеджера для просмотра данных."""

    def __init__(self):  # Конструктор класса
        super().__init__()  # Вызываем конструктор родительского класса
        self.setWindowTitle("Менеджер - Просмотр данных")  # Устанавливаем заголовок окна
        self.setGeometry(100, 100, 600, 400)  # Устанавливаем размеры и положение окна

        # Стилизация окна
        self.setStyleSheet(self.get_style_sheet())  # Устанавливаем стиль через метод

        layout = QVBoxLayout()  # Создаем вертикальную компоновку

        # Заголовок панели менеджера
        self.title_label = QLabel("Панель менеджера")  # Создаем метку заголовка
        self.title_label.setAlignment(Qt.AlignmentFlag.AlignCenter)  # Центрируем заголовок
        self.title_label.setStyleSheet(
            "font-size: 18px; font-weight: bold; margin-bottom: 20px;")  # Стили для заголовка
        layout.addWidget(self.title_label)  # Добавляем заголовок в компоновку

        # Кнопки для просмотра данных
        self.clients_button = QPushButton("Просмотр клиентов")  # Кнопка для просмотра клиентов
        self.clients_button.clicked.connect(self.view_clients)  # Подключаем обработчик
        layout.addWidget(self.clients_button)  # Добавляем кнопку в компоновку

        self.cars_button = QPushButton("Просмотр автомобилей")  # Кнопка для просмотра автомобилей
        self.cars_button.clicked.connect(self.view_cars)  # Подключаем обработчик
        layout.addWidget(self.cars_button)  # Добавляем кнопку в компоновку

        self.orders_button = QPushButton("Просмотр заказов")  # Кнопка для просмотра заказов
        self.orders_button.clicked.connect(self.view_orders)  # Подключаем обработчик
        layout.addWidget(self.orders_button)  # Добавляем кнопку в компоновку

        # Кнопки для экспорта данных
        self.export_clients_button = QPushButton("Экспорт клиентов в Word")  # Кнопка для экспорта клиентов
        self.export_clients_button.clicked.connect(self.export_clients_to_word)  # Подключаем обработчик
        layout.addWidget(self.export_clients_button)  # Добавляем кнопку в компоновку

        self.export_cars_button = QPushButton("Экспорт автомобилей в Word")  # Кнопка для экспорта автомобилей
        self.export_cars_button.clicked.connect(self.export_cars_to_word)  # Подключаем обработчик
        layout.addWidget(self.export_cars_button)  # Добавляем кнопку в компоновку

        self.export_orders_button = QPushButton("Экспорт заказов в Word")  # Кнопка для экспорта заказов
        self.export_orders_button.clicked.connect(self.export_orders_to_word)  # Подключаем обработчик
        layout.addWidget(self.export_orders_button)  # Добавляем кнопку в компоновку

        self.setLayout(layout)  # Устанавливаем компоновку для окна

    def get_style_sheet(self):  # Метод для получения стиля окна
        return """  
            QWidget {
                background-color: #D6EAF8;  
                color: #2C3E50;  
                font-family: Arial;  
                font-size: 14px;  
            }
            QPushButton {
                background-color: #3498DB;  
                border: none;  
                border-radius: 5px;  
                padding: 10px 15px;  
                color: #ECF0F1; 
                font-weight: bold;  
            }
            QPushButton:hover {
                background-color: #2980B9;  
            }
            QPushButton:pressed {
                background-color: #1C5980;  
            }
        """

    def view_clients(self):  # Метод для просмотра клиентов
        self.display_table("clients",
                           ["ID", "Имя", "Фамилия", "Телефон", "Email"])  # Вызываем метод отображения таблицы

    def view_cars(self):  # Метод для просмотра автомобилей
        self.display_table("cars",
                           ["ID", "ID клиента", "Марка", "Модель", "Номер авто"])  # Вызываем метод отображения таблицы

    def view_orders(self):  # Метод для просмотра заказов
        self.display_table("orders", ["ID", "ID авто", "ID пользователя", "Описание", "Цена",
                                      "Дата"])  # Вызываем метод отображения таблицы

    def display_table(self, table_name, headers):  # Метод для отображения данных в таблице
        """Отображает таблицу из базы данных."""
        try:
            conn = sqlite3.connect('myydatabase.db')  # Подключаемся к базе данных
            cursor = conn.cursor()  # Создаем курсор

            # Выполняем запрос для получения всех данных из таблицы
            cursor.execute(f"SELECT * FROM {table_name}")  # Выполняем SQL-запрос
            rows = cursor.fetchall()  # Получаем все строки в результате

            if not rows:  # Если таблица пустая
                QMessageBox.information(self, "Информация", f"Таблица '{table_name}' пуста.")  # Сообщение об информации
                return

            # Создаем виджет таблицы
            table_widget = QTableWidget()  # Создаем экземпляр QTableWidget
            table_widget.setColumnCount(len(headers))  # Устанавливаем число колонок
            table_widget.setHorizontalHeaderLabels(headers)  # Устанавливаем заголовки колонок
            table_widget.setRowCount(len(rows))  # Устанавливаем количество строк

            # Заполняем таблицу данными
            for row_idx, row_data in enumerate(rows):  # Проходим по каждой строке
                for col_idx, col_data in enumerate(row_data):  # Проходим по каждому элементу в строке
                    table_widget.setItem(row_idx, col_idx, QTableWidgetItem(str(col_data)))  # Заполняем таблицу данными

            table_widget.resize(525, 600)  # Устанавливаем размеры таблицы
            table_widget.setWindowTitle(f"Просмотр таблицы {table_name}")  # Устанавливаем заголовок окна таблицы
            table_widget.show()  # Показываем таблицу

            # Сохраняем таблицу как атрибут, чтобы окно не закрывалось
            self.current_table = table_widget  # Сохраняем ссылку на таблицу

        except sqlite3.Error as e:  # Обрабатываем ошибки при работе с базой данных
            QMessageBox.critical(self, "Ошибка", f"Ошибка при работе с базой данных: {e}")  # Сообщение об ошибке
        finally:
            conn.close()  # Закрываем соединение

    def export_clients_to_word(self):  # Метод для экспорта клиентов в Word
        """Экспортирует данные клиентов в Word."""
        self.export_to_word("clients", ["ID", "Имя", "Фамилия", "Телефон", "Email"])  # Вызываем метод экспорта

    def export_cars_to_word(self):  # Метод для экспорта автомобилей в Word
        """Экспортирует данные автомобилей в Word."""
        self.export_to_word("cars", ["ID", "ID клиента", "Марка", "Модель", "Номер авто"])  # Вызываем метод экспорта

    def export_orders_to_word(self):  # Метод для экспорта заказов в Word
        """Экспортирует данные заказов в Word."""
        self.export_to_word("orders",
                            ["ID", "ID авто", "ID пользователя", "Описание", "Цена", "Дата"])  # Вызываем метод экспорта

    def export_to_word(self, table_name, headers):  # Метод для экспорта данных в Word
        """Экспортирует данные из таблицы в документ Word."""
        try:
            conn = sqlite3.connect('myydatabase.db')  # Подключаемся к базе данных
            cursor = conn.cursor()  # Создаем курсор
            cursor.execute(f"SELECT * FROM {table_name}")  # Выполняем SQL-запрос
            data = cursor.fetchall()  # Получаем данные

            if not data:  # Если нет данных для экспорта
                QMessageBox.warning(self, "Ошибка",
                                    f"Нет данных для экспорта из таблицы '{table_name}'.")  # Сообщение об ошибке
                return

            # Создаем документ Word
            doc = Document()  # Создаем экземпляр Document
            doc.add_heading(f'Данные из таблицы {table_name}', level=1)  # Добавляем заголовок в документ
            # Добавление заголовков
            doc.add_paragraph(", ".join(headers))  # Добавляем заголовки в виде строки

            # Добавляем строки из базы данных в документ
            for row in data:  # Проходим по каждой строке данных
                doc.add_paragraph(", ".join(map(str, row)))  # Добавляем строку в документ

            # Сохранение файла
            file_name, _ = QFileDialog.getSaveFileName(self, "Сохранить файл как", "",
                                                       "Word Files (*.docx);;All Files (*)")  # Открываем диалог для сохранения файла
            if file_name:  # Если имя файла задано
                doc.save(file_name)  # Сохраняем документ
                QMessageBox.information(self, "Успех",
                                        f"Данные успешно экспортированы в {file_name}!")  # Сообщение об успешном экспорте

        except Exception as e:  # Обработка ошибок
            QMessageBox.critical(self, "Ошибка", f"Ошибка при экспорте в Word: {e}")  # Сообщение об ошибке

        finally:
            conn.close()  # Закрываем соединение


if __name__ == "__main__":  # Проверяем, является ли файл исполняемым
    app = QtWidgets.QApplication(sys.argv)  # Создаем экземпляр приложения
    auth_app = AuthApp()  # Создаем экземпляр приложения авторизации
    auth_app.show()  # Показываем окно приложения авторизации
    sys.exit(app.exec())  # Запускаем главный цикл приложения
